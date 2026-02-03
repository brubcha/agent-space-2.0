"""
Marketing Kit DOCX Generator  —  FIXED

Root causes found in the Feb 2 output and fixed here:

  BUG 1 — 81 duplicated labels
      The LLM returns e.g. "Profile:Profile: value".
      FIX: _dedup() scrubs back-to-back repeated substrings first.

  BUG 2 — 3 raw # headings and 6 raw **bold** blocks dumped as Normal
      FIX: _parse() detects them and emits H3 blocks.

  BUG 3 — 21 persona label fields rendered as plain Normal text
      FIX: _parse() detects "Label: value" lines and emits LBL blocks.
           _write_blocks() renders them as bold label + normal value.

  BUG 4 — duplicate H1 "Campaign Structure"
      FIX: _seen_h1 set — second write of the same H1 is silently skipped.

  BUG 5 — flat structure (7 H3 vs Swift's 52)
      Sub-headings like "Macro Trends & Growth", "Evergreen Campaigns",
      "Brand Essence", "Landing Page Structure" etc. were never styled.
      FIX: _SUB_HEADINGS whitelist catches known sub-headings precisely.
           A short-line heuristic catches anything else (2-7 words,
           title-case, no period, followed by body text).
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re


# ============================================================================
# DEDUPLICATION
# ============================================================================
_DUPE_RE = re.compile(r'(.{4,}?)\1+')

def _dedup(text: str) -> str:
    return _DUPE_RE.sub(r'\1', text)


# ============================================================================
# BLOCK TYPES
# ============================================================================
H3   = "h3"
BULL = "bullet"
NUM  = "number"
LBL  = "label"
NORM = "normal"

# ── Known persona / structured-field labels ────────────────────────────────
_LABEL_WORDS = {
    "Profile", "Motivation", "Needs", "Messaging", "Demographic",
    "Psychographic", "Buying Behavior", "Summary", "Copy", "Hashtags",
    "Design Goal", "Frequency", "Alignment", "Actions", "Key Components",
    "Focus Areas", "Key Metrics", "Purpose", "Design Standards",
    "Copy Structure", "Posting Cadence", "Content Sources", "Design Tips",
    "Copy Guidelines", "Idea Starters", "Cadence and Governance",
}

# ── Known sub-heading phrases that should always become H3 ─────────────────
# (exact-match after strip; case-insensitive comparison)
_SUB_HEADINGS = {
    "macro trends & growth",
    "competitor landscape & buying behavior",
    "channel opportunities",
    "brand essence",
    "brand purpose",
    "brand personality",
    "tone & voice examples",
    "voice in action",
    "factual foundations",
    "taglines (evaluated)",
    "client do's & don'ts",
    "do's & don'ts",
    "content preferences",
    "required content mix",
    "creative emphases",
    "primary goals",
    "primary post types",
    "post examples",
    "social examples",
    "social production checklist",
    "content guidelines",
    "campaign framework",
    "campaign types",
    "evergreen campaigns",
    "prospecting campaigns",
    "event/launch campaigns",
    "landing page structure",
    "landing page types",
    "foundation priorities",
    "measurement framework",
    "keyword strategy",
    "keyword opportunity analysis",
    "blog strategy",
    "blog structure",
    "blog structure guidelines",
    "content execution framework",
    "static feed",
    "dynamic stories",
    "dynamic reels",
    "dynamic videos",
    "static ugc",
    "community spotlight",
    "health assessment landing page",
    "specific condition landing page",
    "comparison landing page",
}


# ============================================================================
# PARSER
# ============================================================================
def _parse(raw: str) -> list:
    if not raw or not raw.strip():
        return []

    blocks = []
    chunks = re.split(r'\n\s*\n', raw.strip())

    for chunk in chunks:
        lines = chunk.split('\n')
        i = 0
        while i < len(lines):
            line = _dedup(lines[i]).strip()
            if not line:
                i += 1
                continue

            # ── 1. # or ## heading ──────────────────────────────────
            m = re.fullmatch(r'#{1,4}\s+(.+)', line)
            if m:
                blocks.append({"type": H3, "text": m.group(1).strip()})
                i += 1
                continue

            # ── 2. **Heading** alone on a line ──────────────────────
            m = re.fullmatch(r'\*\*(.+?)\*\*:?\s*$', line)
            if m:
                blocks.append({"type": H3, "text": m.group(1).strip()})
                i += 1
                continue

            # ── 3. 0X | Finding Title ───────────────────────────────
            m = re.fullmatch(r'\*?\*?(\d{1,2})\s*[|]\s*(.+?)\*?\*?\s*$', line)
            if m:
                blocks.append({"type": H3,
                               "text": f"{m.group(1).zfill(2)} | {m.group(2).strip()}"})
                i += 1
                continue

            # ── 4. bullet  •  or  -  ────────────────────────────────
            m = re.match(r'[•\-]\s+(.*)', line)
            if m:
                blocks.append({"type": BULL, "text": _dedup(m.group(1)).strip()})
                i += 1
                continue

            # ── 5. numbered  1.  ────────────────────────────────────
            m = re.match(r'\d+\.\s+(.*)', line)
            if m:
                blocks.append({"type": NUM, "text": _dedup(m.group(1)).strip()})
                i += 1
                continue

            # ── 6. Label: value  ────────────────────────────────────
            m = re.match(r'\*{0,2}([A-Za-z &\'&()]+?):\*{0,2}\s+(.*)', line)
            if m and m.group(1).strip() in _LABEL_WORDS:
                blocks.append({"type": LBL,
                               "label": m.group(1).strip(),
                               "text":  _dedup(m.group(2)).strip()})
                i += 1
                continue

            # ── 7. "The Persona Name:" line ─────────────────────────
            m = re.match(r'(The [A-Z][^:]{4,40}):\s*(.*)', line)
            if m:
                blocks.append({"type": H3, "text": m.group(1).strip()})
                remainder = _dedup(m.group(2)).strip()
                if remainder:
                    blocks.append({"type": NORM, "text": remainder})
                i += 1
                continue

            # ── 8. Known sub-heading (whitelist) ────────────────────
            if line.lower().rstrip(':') in _SUB_HEADINGS:
                blocks.append({"type": H3, "text": line.rstrip(':')})
                i += 1
                continue

            # ── 9. Short standalone line heuristic ──────────────────
            # 2-7 words, starts with capital, no trailing period or colon,
            # not a checkbox line, not starting with a digit, and there IS
            # a following non-blank line (so it's not a dangling fragment).
            words_here = line.split()
            if (2 <= len(words_here) <= 7 and
                not line.endswith('.') and
                not line.endswith(':') and
                words_here[0][0].isupper() and
                not line.startswith('✅') and
                not line.startswith('❌') and
                not re.match(r'\d', line) and
                not line.startswith('"') and
                not line.startswith("'")):
                # peek forward for a following line
                has_next = any(lines[j].strip() for j in range(i+1, len(lines)))
                if has_next:
                    blocks.append({"type": H3, "text": line})
                    i += 1
                    continue

            # ── 10. normal text ─────────────────────────────────────
            collected = [line]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt:
                    break
                # stop at any special-pattern line
                if (re.fullmatch(r'#{1,4}\s+.+', nxt) or
                    re.fullmatch(r'\*\*(.+?)\*\*:?\s*$', nxt) or
                    re.fullmatch(r'\*?\*?\d{1,2}\s*[|]\s*.+?\*?\*?\s*$', nxt) or
                    re.match(r'[•\-]\s', nxt) or
                    re.match(r'\d+\.\s', nxt) or
                    re.match(r'The [A-Z][^:]{4,40}:', nxt) or
                    nxt.lower().rstrip(':') in _SUB_HEADINGS):
                    break
                # also stop if next line would hit the short-line heuristic
                nxt_words = nxt.split()
                if (2 <= len(nxt_words) <= 7 and
                    not nxt.endswith('.') and not nxt.endswith(':') and
                    nxt_words[0][0].isupper() and
                    not nxt.startswith('✅') and not nxt.startswith('❌') and
                    not re.match(r'\d', nxt) and
                    not nxt.startswith('"') and not nxt.startswith("'")):
                    # check if there's a line after *that*
                    if any(lines[j].strip() for j in range(i+1, len(lines))):
                        break
                collected.append(_dedup(nxt))
                i += 1
            blocks.append({"type": NORM, "text": " ".join(collected)})

    return blocks


# ============================================================================
# BLOCK WRITER
# ============================================================================
def _write_blocks(doc: Document, blocks: list):
    # Remove consecutive duplicate H3s (artifact of LLM repeating sub-headings)
    cleaned = []
    prev_h3 = None
    for b in blocks:
        if b["type"] == "h3":
            if b["text"] == prev_h3:
                continue
            prev_h3 = b["text"]
        else:
            prev_h3 = None
        cleaned.append(b)
    blocks = cleaned
    for b in blocks:
        t = b["type"]

        if t == H3:
            doc.add_heading(b["text"], level=3)

        elif t == BULL:
            doc.add_paragraph(b["text"], style="List Bullet")

        elif t == NUM:
            doc.add_paragraph(b["text"], style="List Number")

        elif t == LBL:
            p = doc.add_paragraph()
            run_l = p.add_run(f"{b['label']}:")
            run_l.bold = True
            if b.get("text"):
                p.add_run(f"  {b['text']}")

        else:   # NORM
            raw = b["text"]
            parts = re.split(r'\*\*(.+?)\*\*', raw)
            if len(parts) == 1:
                doc.add_paragraph(raw)
            else:
                p = doc.add_paragraph()
                for idx, part in enumerate(parts):
                    if not part:
                        continue
                    run = p.add_run(part)
                    if idx % 2 == 1:
                        run.bold = True


# ============================================================================
# GENERATOR CLASS
# ============================================================================
class MarketingKitDocxGenerator:
    def __init__(self, design_system=None):
        self.doc = Document()
        self.design_system = design_system
        self._seen_h1 = set()
        self._setup_styles()
        self._setup_margins()

    def _setup_margins(self):
        section = self.doc.sections[0]
        section.top_margin    = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin   = Inches(0.5)
        section.right_margin  = Inches(0.5)

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font  = style.font
        if self.design_system and hasattr(self.design_system, 'typography'):
            font.name = self.design_system.typography.body_font
            font.size = Pt(self.design_system.typography.body_size)
        else:
            font.name = 'Arial'
            font.size = Pt(11)

    def _add_title(self, text: str):
        heading = self.doc.add_heading(text, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.size = Pt(24)
        run.font.bold = True
        if self.design_system and hasattr(self.design_system, 'colors'):
            run.font.color.rgb = RGBColor(*self.design_system.colors.primary)

    def _add_h1(self, text: str):
        if text in self._seen_h1:
            return
        self._seen_h1.add(text)
        heading = self.doc.add_heading(text, level=1)
        run = heading.runs[0]
        run.font.size  = Pt(18)
        run.font.bold  = True
        if self.design_system and hasattr(self.design_system, 'colors'):
            run.font.color.rgb = RGBColor(*self.design_system.colors.secondary)

    def _add_paragraph(self, text: str, bold: bool = False):
        text = _dedup(text)
        p = self.doc.add_paragraph(text)
        if bold and p.runs:
            p.runs[0].bold = True
        return p

    def _write_section(self, content: str):
        _write_blocks(self.doc, _parse(str(content)))

    def generate(self, company_name: str, agent_results: dict, save_path: str):
        # title page
        self._add_title("Marketing Kit")
        self.doc.add_paragraph()
        tp = self.doc.add_paragraph(company_name)
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.runs[0].font.size = Pt(16)
        self.doc.add_paragraph()
        dp = self.doc.add_paragraph(datetime.now().strftime("%B %Y"))
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.runs[0].font.size = Pt(12)
        dp.runs[0].italic    = True
        self.doc.add_page_break()

        # TOC
        self._add_h1("Table of Contents")
        for i, title in enumerate([
            "Overview", "The Goal", "Key Findings", "Market Landscape",
            "Audience & Personas", "Brand Voice", "Content Strategy",
            "Social Strategy", "Campaign Structure", "Engagement Framework"
        ], 1):
            self._add_paragraph(f"{i}. {title}")
        self.doc.add_page_break()

        # sections
        SECTIONS = [
            ("overview_writer",              "Overview"),
            ("key_findings_researcher",      "Key Findings"),
            ("market_landscape_analyzer",    "Market Landscape"),
            ("persona_creator",              "Audience & Personas"),
            ("brand_voice_definer",          "Brand Voice"),
            ("keyword_strategist",           "Keyword Strategy"),
            ("social_strategist",            "Social Strategy"),
            ("campaign_architect",           "Campaign Structure"),
            ("engagement_framework_builder", "Engagement Framework"),
        ]

        for agent_key, section_title in SECTIONS:
            if agent_key not in agent_results:
                continue
            output = agent_results[agent_key].get("output", "")
            if not output or not str(output).strip():
                continue
            self._add_h1(section_title)
            self._write_section(output)
            self.doc.add_page_break()

        self.doc.save(save_path)
        print(f"✓ Marketing kit saved to: {save_path}")


# ============================================================================
# CONVENIENCE FUNCTION
# ============================================================================
def generate_marketing_kit_docx(company_name: str, agent_results: dict,
                                output_dir: str = ".", design_system=None):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"Marketing_Kit_{company_name.replace(' ', '_')}_{timestamp}.docx"
    save_path = f"{output_dir}/{filename}"
    generator = MarketingKitDocxGenerator(design_system=design_system)
    generator.generate(company_name, agent_results, save_path)
    return save_path
